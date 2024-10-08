import streamlit as st
from googleapiclient.discovery import build
from datetime import datetime
import os
import openai
import json
from docx import Document
from io import BytesIO
from dotenv import load_dotenv
from collections import defaultdict

load_dotenv()

# Get API keys from environment variables
YOUTUBE_API_KEY = os.environ.get("YOUTUBE_API_KEY")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

if not YOUTUBE_API_KEY:
    st.error("YouTube API key not found. Please set the YOUTUBE_API_KEY environment variable.")
    st.stop()

if not OPENAI_API_KEY:
    st.error("OpenAI API key not found. Please set the OPENAI_API_KEY environment variable.")
    st.stop()

youtube = build('youtube', 'v3', developerKey=YOUTUBE_API_KEY)
openai.api_key = OPENAI_API_KEY

def get_all_comments(video_id: str):
    try:
        comments = []
        nextPageToken = None
        while True:
            response = youtube.commentThreads().list(
                part='snippet',
                videoId=video_id,
                maxResults=100,  # Maximum allowed by the API
                pageToken=nextPageToken,
                order='time'
            ).execute()

            for item in response['items']:
                comment = item['snippet']['topLevelComment']['snippet']
                comments.append({
                    'author': comment['authorDisplayName'],
                    'text': comment['textDisplay'],
                    'likes': comment['likeCount'],
                    'published_at': datetime.strptime(comment['publishedAt'], "%Y-%m-%dT%H:%M:%SZ")
                })

                # Fetch replies to this comment
                if item['snippet']['totalReplyCount'] > 0:
                    replies = youtube.comments().list(
                        part='snippet',
                        parentId=item['id'],
                        maxResults=100  # Maximum allowed by the API
                    ).execute()
                    
                    for reply in replies['items']:
                        reply_snippet = reply['snippet']
                        comments.append({
                            'author': reply_snippet['authorDisplayName'],
                            'text': reply_snippet['textDisplay'],
                            'likes': reply_snippet['likeCount'],
                            'published_at': datetime.strptime(reply_snippet['publishedAt'], "%Y-%m-%dT%H:%M:%SZ")
                        })

            nextPageToken = response.get('nextPageToken')
            if not nextPageToken:
                break

        return comments
    except Exception as e:
        return f"An error occurred while fetching comments: {str(e)}"

def extract_questions(comments, video_info):
    try:
        comments_with_authors = [f"{comment['author']}: {comment['text']} (Date: {comment['published_at'].strftime('%Y-%m-%d %H:%M:%S')})" for comment in comments]
        all_comments_text = "\n".join(comments_with_authors[:100])  # Limit to first 100 comments to avoid token limit
        
        prompt = f"""Analyze the following YouTube comments for the video titled "{video_info['title']}" and extract the 4 most relevant direct questions and 4 most relevant indirect questions about the video content. Improve and rephrase the questions to make them more efficient, clear, and insightful.

Video Title: {video_info['title']}
Video Description: {video_info['description'][:500]}  # Limit description length

Comments:
{all_comments_text}

Please follow these guidelines:
1. Provide exactly 4 direct questions and 4 indirect questions.
2. Ensure all questions are directly relevant to the video content.
3. For each question, provide the commenter's name and the date and time of the comment.
4. Improve and rephrase each question to make it more clear, concise, and insightful.
5. Do not add any additional context or explanations to the questions.

Format your response as follows:
Direct Questions:
1. [Improved Direct Question 1] (Commenter: [Name], Date: [YYYY-MM-DD HH:MM:SS])
2. [Improved Direct Question 2] (Commenter: [Name], Date: [YYYY-MM-DD HH:MM:SS])
3. [Improved Direct Question 3] (Commenter: [Name], Date: [YYYY-MM-DD HH:MM:SS])
4. [Improved Direct Question 4] (Commenter: [Name], Date: [YYYY-MM-DD HH:MM:SS])

Indirect Questions:
1. [Improved Indirect Question 1] (Commenter: [Name], Date: [YYYY-MM-DD HH:MM:SS])
2. [Improved Indirect Question 2] (Commenter: [Name], Date: [YYYY-MM-DD HH:MM:SS])
3. [Improved Indirect Question 3] (Commenter: [Name], Date: [YYYY-MM-DD HH:MM:SS])
4. [Improved Indirect Question 4] (Commenter: [Name], Date: [YYYY-MM-DD HH:MM:SS])

If there are not enough relevant questions in either category, write 'No more relevant questions found.' for the remaining slots.
"""

        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an AI assistant specialized in analyzing YouTube comments and extracting insightful, relevant questions. Your task is to identify, categorize, improve, and limit the number of questions from user comments, ensuring they are directly related to the video content."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000
        )

        return response.choices[0].message.content
    except openai.error.InvalidRequestError as e:
        st.error(f"Error in OpenAI API request: {str(e)}")
        return "Unable to extract questions due to an API error. This may be due to the length of the input. Try analyzing a video with fewer comments."
    except Exception as e:
        st.error(f"An unexpected error occurred: {str(e)}")
        return "An unexpected error occurred while extracting questions."

def generate_related_questions(questions):
    try:
        prompt = f"""Based on the following extracted questions from YouTube comments, generate a list of 5-10 related questions that could further enhance the discussion about the video content. These related questions should explore themes or topics that are implied by the original questions but not directly asked.

Extracted Questions:
{questions}

Please provide a list of related questions that:
1. Expand on the themes present in the original questions
2. Explore potential implications or consequences related to the topics discussed
3. Encourage deeper analysis or critical thinking about the video content
4. Address potential gaps in the discussion that the original questions might have missed

Format your response as a numbered list of questions.
"""

        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an AI assistant specialized in generating insightful and related questions based on existing questions from YouTube comments."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000
        )

        return response.choices[0].message.content
    except Exception as e:
        st.error(f"An error occurred while generating related questions: {str(e)}")
        return "Unable to generate related questions due to an error."

def get_video_info(video_id):
    try:
        response = youtube.videos().list(
            part='snippet,statistics',
            id=video_id
        ).execute()

        if 'items' in response:
            video = response['items'][0]
            return {
                'title': video['snippet']['title'],
                'description': video['snippet']['description'],
                'views': video['statistics']['viewCount'],
                'likes': video['statistics']['likeCount'],
                'comments': video['statistics']['commentCount'],
                'published_at': video['snippet']['publishedAt'],
                'thumbnail': video['snippet']['thumbnails']['high']['url']
            }
        else:
            return None
    except Exception as e:
        st.error(f"An error occurred while fetching video info: {str(e)}")
        return None

def analyze_comment_sentiment(comment):
    try:
        prompt = f"Analyze the sentiment of the following comment and classify it as POSITIVE, NEGATIVE, or NEUTRAL. Respond with only one word: POSITIVE, NEGATIVE, or NEUTRAL.\n\nComment: {comment['text']}"
        
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a sentiment analysis AI. Classify the given comment as POSITIVE, NEGATIVE, or NEUTRAL."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=10
        )
        
        sentiment = response.choices[0].message.content.strip().upper()
        return sentiment if sentiment in ["POSITIVE", "NEGATIVE", "NEUTRAL"] else "NEUTRAL"
    except Exception as e:
        st.warning(f"Error analyzing sentiment for a comment: {str(e)}")
        return "NEUTRAL"

def analyze_comments(video_id):
    if video_id:
        with st.spinner("📊 Fetching and analyzing comments..."):
            comments = get_all_comments(video_id)
            if isinstance(comments, list):
                st.session_state.comments = comments
                st.session_state.comments.sort(key=lambda x: x['published_at'], reverse=True)
                st.session_state.video_info = get_video_info(video_id)
                
                if st.session_state.video_info:
                    st.session_state.questions = extract_questions(comments[:100], st.session_state.video_info)  # Limit to 100 comments
                    if st.session_state.questions and not st.session_state.questions.startswith("Unable to extract questions"):
                        st.session_state.related_questions = generate_related_questions(st.session_state.questions)
                    else:
                        st.session_state.related_questions = "Unable to generate related questions due to an error in extracting initial questions."
                else:
                    st.error("Unable to fetch video information. Please check the video ID and try again.")
                    return

                # Analyze sentiment for each comment
                sentiments = defaultdict(list)
                for comment in comments:
                    sentiment = analyze_comment_sentiment(comment)
                    comment['sentiment'] = sentiment
                    sentiments[sentiment].append(comment)
                
                st.session_state.sentiments = dict(sentiments)
                st.session_state.sentiment_counts = {
                    'POSITIVE': len(sentiments['POSITIVE']),
                    'NEGATIVE': len(sentiments['NEGATIVE']),
                    'NEUTRAL': len(sentiments['NEUTRAL'])
                }
                st.session_state.total_comments_analyzed = len(comments)
            else:
                st.error(comments)
    else:
        st.error("⚠️ Please enter a YouTube Video ID.")

def create_docx_report(video_info, comments, questions, related_questions, sentiment_counts, sentiments):
    doc = Document()
    doc.add_heading('YouTube Video Analysis Report', 0)

    # Video Information
    doc.add_heading('Video Information', level=1)
    doc.add_paragraph(f"Title: {video_info['title']}")
    doc.add_paragraph(f"Views: {video_info['views']}")
    doc.add_paragraph(f"Likes: {video_info['likes']}")
    doc.add_paragraph(f"Comments: {video_info['comments']}")
    doc.add_paragraph(f"Published: {video_info['published_at']}")

    # Sentiment Analysis
    doc.add_heading('Sentiment Analysis', level=1)
    total_comments = sum(sentiment_counts.values())
    doc.add_paragraph(f"Total comments analyzed: {total_comments}")
    doc.add_paragraph(f"Positive: {sentiment_counts['POSITIVE']} ({(sentiment_counts['POSITIVE'] / total_comments) * 100:.2f}%)")
    doc.add_paragraph(f"Neutral: {sentiment_counts['NEUTRAL']} ({(sentiment_counts['NEUTRAL'] / total_comments) * 100:.2f}%)")
    doc.add_paragraph(f"Negative: {sentiment_counts['NEGATIVE']} ({(sentiment_counts['NEGATIVE'] / total_comments) * 100:.2f}%)")

    # Add sentiment-specific comments
    for sentiment in ['POSITIVE', 'NEUTRAL', 'NEGATIVE']:
        doc.add_heading(f"{sentiment.capitalize()} Comments", level=2)
        for comment in sentiments[sentiment]:
            doc.add_paragraph(f"Author: {comment['author']}")
            doc.add_paragraph(f"Text: {comment['text']}")
            doc.add_paragraph(f"Likes: {comment['likes']}")
            doc.add_paragraph(f"Published at: {comment['published_at']}")
            doc.add_paragraph("---")

    # Extracted and Improved Questions
    doc.add_heading('Extracted and Improved Questions', level=1)
    doc.add_paragraph(questions)

    # Related Questions
    doc.add_heading('Related Questions', level=1)
    doc.add_paragraph(related_questions)

    return doc

st.set_page_config(layout="wide", page_title="Bent's Comment Analyzer 🎥💬")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600&display=swap');
    
    body {
        font-family: 'Poppins', sans-serif;
        background-color: #f0f2f5;
        color: #1a1a1a;
    }
    .main {
        max-width: 1200px;
        margin: 0 auto;
        background-color: #ffffff;
        padding: 2rem;
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    h1 {
        color: #3498db;
        text-align: center;
        font-size: 2.5em;
        margin-bottom: 1em;
    }
    h2 {
        color: #2c3e50;
        font-size: 1.8em;
        margin-top: 1em;
        margin-bottom: 0.5em;
    }
    .stTextInput>div>div>input {
        background-color: #f1f3f5;
        color: #333;
        border: 2px solid #3498db;
        border-radius: 5px;
        padding: 10px 15px;
    }
    .stButton>button {
        background-color: #3498db;
        color: white;
        border: none;
        padding: 10px 20px;
        text-align: center;
        text-decoration: none;
        display: inline-block;
        font-size: 16px;
        margin: 4px 2px;
        cursor: pointer;
        border-radius: 5px;
        transition: background-color 0.3s, color 0.3s;
    }
    .stButton>button:hover {
        background-color: #2980b9;
        color: #ffffff;
    }
    .small-button {
        padding: 5px 10px;
        font-size: 14px;
    }
    .comment {
        background-color: #f9f9f9;
        border-left: 4px solid #3498db;
        padding: 15px;
        margin-bottom: 15px;
        border-radius: 4px;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
    }
    .comment-author {
        font-weight: bold;
        color: #2c3e50;
    }
    .comment-date {
        font-size: 0.8em;
        color: #7f8c8d;
    }
    .comment-text {
        margin-top: 5px;
        color: #34495e;
    }
    .comment-likes {
        font-size: 0.9em;
        color: #e74c3c;
        margin-top: 5px;
    }
    .sort-button {
        background-color: #2ecc71;
        margin-bottom: 10px;
    }
    .sort-button:hover {
        background-color: #27ae60;
    }
    .summary-box {
        background-color: #ecf0f1;
        border-radius: 5px;
        padding: 10px;
        margin-bottom: 15px;
        font-size: 0.9em;
    }
</style>
""", unsafe_allow_html=True)

st.title("🎥 Bent's Comment Analyzer 💬")

video_id = st.text_input("🔍 Enter YouTube Video ID")

if 'comments' not in st.session_state:
    st.session_state.comments = []
if 'sort_order' not in st.session_state:
    st.session_state.sort_order = 'newest'
if 'show_comments' not in st.session_state:
    st.session_state.show_comments = 10
if 'questions' not in st.session_state:
    st.session_state.questions = None
if 'related_questions' not in st.session_state:
    st.session_state.related_questions = None
if 'video_info' not in st.session_state:
    st.session_state.video_info = None
if 'sentiments' not in st.session_state:
    st.session_state.sentiments = None
if 'sentiment_counts' not in st.session_state:
    st.session_state.sentiment_counts = None
if 'total_comments_analyzed' not in st.session_state:
    st.session_state.total_comments_analyzed = 0

def toggle_sort_order():
    st.session_state.sort_order = 'oldest' if st.session_state.sort_order == 'newest' else 'newest'
    st.session_state.comments.sort(key=lambda x: x['published_at'], reverse=(st.session_state.sort_order == 'newest'))

def show_more_comments():
    st.session_state.show_comments = min(st.session_state.show_comments + 10, len(st.session_state.comments))

def show_less_comments():
    st.session_state.show_comments = max(st.session_state.show_comments - 10, 10)

if st.button("🚀 Analyze Comments", key="analyze_button"):
    analyze_comments(video_id)

# Display video information, sentiment analysis, and export data at the top
if st.session_state.video_info:
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("## 📺 Video Information")
        st.markdown(f"**Title:** {st.session_state.video_info['title']}")
        st.markdown(f"**Views:** {st.session_state.video_info['views']}")
        st.markdown(f"**Likes:** {st.session_state.video_info['likes']}")
        st.markdown(f"**Comments:** {st.session_state.video_info['comments']}")
        st.markdown(f"**Published:** {st.session_state.video_info['published_at']}")
    
    with col2:
        st.image(st.session_state.video_info['thumbnail'], use_column_width=True)

if 'sentiment_counts' in st.session_state and st.session_state.sentiment_counts:
    st.markdown("## 💭 Sentiment Analysis")
    
    # Add summary box
    st.markdown(f"""
    <div class="summary-box">
        Total comments analyzed: {st.session_state.total_comments_analyzed}<br>
        Positive: {st.session_state.sentiment_counts['POSITIVE']}<br>
        Neutral: {st.session_state.sentiment_counts['NEUTRAL']}<br>
        Negative: {st.session_state.sentiment_counts['NEGATIVE']}
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    total_comments = sum(st.session_state.sentiment_counts.values())
    
    def show_sentiment_comments(sentiment):
        st.session_state.active_sentiment = sentiment
    
    with col1:
        positive_percentage = (st.session_state.sentiment_counts.get('POSITIVE', 0) / total_comments) * 100 if total_comments else 0
        st.metric("Positive", f"{positive_percentage:.2f}%")
        if st.button("Show Positive Comments"):
            show_sentiment_comments('POSITIVE')
    
    with col2:
        neutral_percentage = (st.session_state.sentiment_counts.get('NEUTRAL', 0) / total_comments) * 100 if total_comments else 0
        st.metric("Neutral", f"{neutral_percentage:.2f}%")
        if st.button("Show Neutral Comments"):
            show_sentiment_comments('NEUTRAL')
    
    with col3:
        negative_percentage = (st.session_state.sentiment_counts.get('NEGATIVE', 0) / total_comments) * 100 if total_comments else 0
        st.metric("Negative", f"{negative_percentage:.2f}%")
        if st.button("Show Negative Comments"):
            show_sentiment_comments('NEGATIVE')

    if 'active_sentiment' in st.session_state and 'sentiments' in st.session_state:
        st.markdown(f"## {st.session_state.active_sentiment.capitalize()} Comments")
        for comment in st.session_state.sentiments.get(st.session_state.active_sentiment, []):
            st.markdown(f"""
            <div class="comment">
                <div class="comment-author">👤 {comment['author']}</div>
                <div class="comment-date">🕒 {comment['published_at'].strftime('%Y-%m-%d %H:%M:%S')}</div>
                <div class="comment-text">{comment['text']}</div>
                <div class="comment-likes">❤️ {comment['likes']}</div>
            </div>
            """, unsafe_allow_html=True)
else:
    st.info("No sentiment analysis data available. Please analyze comments first.")

if st.session_state.comments:
    st.markdown("## 📤 Export Data")
    export_format = st.selectbox("Choose export format:", ["CSV", "JSON", "DOCX"])
    
    if st.button("Export Data"):
        if export_format == "CSV":
            csv = "Author,Text,Likes,Published At,Sentiment\n"
            for comment in st.session_state.comments:
                csv += f"{comment['author']},{comment['text'].replace(',', ' ')},{comment['likes']},{comment['published_at']},{comment.get('sentiment', 'NEUTRAL')}\n"
            st.download_button(
                label="Download CSV",
                data=csv,
                file_name="youtube_analysis.csv",
                mime="text/csv"
            )
        elif export_format == "JSON":
            data = {
                "video_info": st.session_state.video_info,
                "sentiment_counts": st.session_state.sentiment_counts,
                "questions": st.session_state.questions,
                "related_questions": st.session_state.related_questions,
                "comments": st.session_state.comments,
                "total_comments_analyzed": st.session_state.total_comments_analyzed
            }
            json_str = json.dumps(data, default=str)
            st.download_button(
                label="Download JSON",
                data=json_str,
                file_name="youtube_analysis.json",
                mime="application/json"
            )
        else:  # DOCX
            doc = create_docx_report(
                st.session_state.video_info,
                st.session_state.comments,
                st.session_state.questions,
                st.session_state.related_questions,
                st.session_state.sentiment_counts,
                st.session_state.sentiments
            )
            bio = BytesIO()
            doc.save(bio)
            st.download_button(
                label="Download DOCX",
                data=bio.getvalue(),
                file_name="youtube_analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Display comments and extracted questions
if st.session_state.comments:
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("<h2>📝 Comments</h2>", unsafe_allow_html=True)
        if st.button(f"{'🔽' if st.session_state.sort_order == 'newest' else '🔼'} Sort: {st.session_state.sort_order.capitalize()}", 
                    key="sort_button", help="Toggle between newest and oldest comments"):
            toggle_sort_order()
        
        for i, comment in enumerate(st.session_state.comments[:st.session_state.show_comments]):
            st.markdown(f"""
            <div class="comment">
                <div class="comment-author">👤 {comment['author']}</div>
                <div class="comment-date">🕒 {comment['published_at'].strftime('%Y-%m-%d %H:%M:%S')}</div>
                <div class="comment-text">{comment['text']}</div>
                <div class="comment-likes">❤️ {comment['likes']}</div>
            </div>
            """, unsafe_allow_html=True)
        
        col1_1, col1_2, col1_3 = st.columns([1,1,2])
        with col1_1:
            if st.session_state.show_comments < len(st.session_state.comments):
                if st.button("📥 Show More", key="show_more"):
                    show_more_comments()
        with col1_2:
            if st.session_state.show_comments > 10:
                if st.button("📤 Show Less", key="show_less"):
                    show_less_comments()
        with col1_3:
            st.write(f"Showing {st.session_state.show_comments} of {len(st.session_state.comments)} comments")
    
    with col2:
        st.markdown("<h2>❓ Extracted and Improved Questions</h2>", unsafe_allow_html=True)
        if st.session_state.questions:
            st.markdown(st.session_state.questions, unsafe_allow_html=True)
        else:
            st.info("🤔 No questions extracted yet. Try analyzing a video with more comments or discussions.")
        
        st.markdown("<h2>🔍 Related Questions</h2>", unsafe_allow_html=True)
        if st.session_state.related_questions:
            st.markdown(st.session_state.related_questions, unsafe_allow_html=True)
        else:
            st.info("💡 No related questions generated yet. Try analyzing a video to get related questions.")

st.markdown("---")
